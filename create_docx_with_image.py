import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_image_with_caption(doc, img_path, caption, width_inches=5.5):
    """Helper to add centered image and caption."""
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
        
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(80, 80, 80)
    else:
        doc.add_paragraph(f"[Image Missing: {img_path}]").font.color.rgb = RGBColor(255, 0, 0)

def main():
    doc = Document()
    
    # --- Header Section ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = header.add_run('Assignment 4: Deep-Dive into LLaVA Vision-Language Architectures\n')
    title.bold = True
    title.font.size = Pt(18)
    
    subtitle = header.add_run('NJIT CS-685: Computer Vision | Spring 2026\n')
    subtitle.font.size = Pt(14)
    
    meta = header.add_run('Author: Saketh Varma Dantuluri | UCID: sd2399')
    meta.italic = True
    meta.font.size = Pt(11)
    
    doc.add_paragraph().add_run('_' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Section 1: Architectural Analysis ---
    doc.add_heading('1. Architectural Analysis [Tasks 1.1 - 1.3]', level=1)
    doc.add_paragraph(
        "The Large Language-and-Vision Assistant (LLaVA) represents an architectural milestone in multimodal AI. "
        "Rather than training a completely new modality-aware model from the ground up, LLaVA employs an 'agentic' "
        "approach, bridging a frozen vision encoder with a powerful large language model (LLM) using a learned projection interface."
    )
    
    add_image_with_caption(doc, "llava_arch.png", "Figure 1: High-level architectural flow of the LLaVA framework.", 5.0)

    doc.add_heading('1.1 The Multimodal Forward Pass Mechanics', level=2)
    doc.add_paragraph(
        "The technical core of LLaVA lies in its ability to treat visual features as 'pseudo-text' tokens. "
        "The process begins with an input image processed through a CLIP ViT-L/14 transformer, "
        "which extracts grid-level visual features. These features are then mapped through a projection matrix (W) "
        "into the LLM's word embedding space, allowing the model to attend to visual concepts as a series of prefix tokens."
    )
    
    doc.add_heading('1.2 The Logic of Minimalist Projection', level=2)
    doc.add_paragraph(
        "A simple projection layer is sufficient because both CLIP and LLaMA are already semantically rich. "
        "The architecture assumes a 'latent space continuity,' where a linear mapping can bridge the gap between "
        "visual and textual manifestations of the same concepts. However, if alignment is poor, the LLM may "
        "ignore visual cues entirely, favoring its own language priors."
    )

    doc.add_heading('1.3 Strategic Integration vs. Native Training', level=2)
    doc.add_paragraph(
        "The decision to use a minimal projection layer preserves the LLM's pre-trained intelligence while dramatically "
        "reducing computational overhead. The primary drawback is the 'Frozen Vision Backbone,' which prevents "
        "the model from learning new visual primitives during multimodal fine-tuning."
    )

    # --- Section 2: Training Pipeline ---
    doc.add_page_break()
    doc.add_heading('2. Quantitative Training Pipeline [Tasks 2.1 - 2.2]', level=1)
    doc.add_paragraph("The pedagogical achievement of LLaVA is its two-stage training regime, which ensures stability while gaining capabilities.")
    
    add_image_with_caption(doc, "training_phases.png", "Figure 2: The progressive transition from freezing to fine-tuning across the pipeline stages.", 5.5)

    doc.add_heading('2.1 The Two-Stage Training Paradigm', level=2)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Stage'
    hdr_cells[1].text = 'Technical Objective'
    hdr_cells[2].text = 'Trained Components'
    hdr_cells[3].text = 'Data Type'
    
    row1 = table.add_row().cells
    row1[0].text = '1: Alignment'
    row1[1].text = 'Sync CLIP to LLM space'
    row1[2].text = 'Projection Layer Only'
    row1[3].text = '595K CC3M pairs'
    
    row2 = table.add_row().cells
    row2[0].text = '2: Tuning'
    row2[1].text = 'Align to human intent'
    row2[2].text = 'LLM + Projection'
    row2[3].text = '158K GPT-4 Dialogs'

    add_image_with_caption(doc, "feature_alignment.png", "Figure 3: Conceptual alignment of visual and textual manifolds.", 5.2)

    doc.add_heading('2.2 Synthetic Data & The GPT-4 Bridge', level=2)
    doc.add_paragraph(
        "LLaVA's success is deeply tied to synthetic multi-turn dialogues generated by language-only GPT-4. "
        "While this allows for massive scaling, the model inherits GPT-4's linguistic biases and risks of 'poisoned' "
        "training data if the original textual descriptions used for generation were incorrect."
    )

    # --- Section 3: Reflection ---
    doc.add_page_break()
    doc.add_heading('3. Comparative Reflection & Synthesis [Tasks 3.1 - 3.3]', level=1)
    
    doc.add_paragraph(
        "The LLaVA project marks a pivotal moment in the transition from specialized vision models to unified multimodal agents. "
        "However, it also highlights the fundamental challenges of 'bolting' together disparate sensory modules that were never designed for joint operation."
    )
    
    doc.add_heading('3.1 Taxonomy of Multimodality: Late-Fusion vs. Native-Fusion', level=2)
    doc.add_paragraph(
        "LLaVA is fundamentally a Language Model conditioned on visual embeddings (late-fusion). "
        "Compared to architectures like BLIP-2 with its Q-Former, LLaVA's simpler projection is more efficient but less selective. "
        "Native multimodal models like Gemini represent 'early-fusion,' where images are processed as fundamental sensory inputs rather than linguistic abstractions."
    )
    
    doc.add_heading('3.2 The Dual-Locus of Alignment', level=2)
    doc.add_paragraph(
        "Alignment is a dual-process phenomenon. Structural alignment (Stage 1) creates the manifold bridge, while Behavioral alignment (Stage 2) "
        "teaches the model to synchronize its intense conversational reasoning with specific visual intents. "
        "Without both, the model suffers from cognitive dissonance between its vision and logic."
    )
    
    doc.add_heading('3.3 The "Frozen Vision" Resolution Ceiling', level=2)
    doc.add_paragraph(
        "The primary bottleneck is the 'frozen vision backbone.' Because CLIP operates at a fixed resolution, fine-grained details "
        "are irretrievably lost before the LLM begins processing. This 'sampling bottleneck' creates a hard ceiling on visual precision "
        "that later models like LLaVA-1.5 attempt to bypass with larger resolution and deeper projection layers."
    )

    # --- Section 4: References ---
    doc.add_heading('4. References & Bibliography', level=1)
    refs = [
        "Liu, H., Li, C., Wu, Q., & Lee, Y. J. (2023). Visual Instruction Tuning. NeurIPS.",
        "Radford, A., et al. (2021). CLIP. arXiv preprint arXiv:2103.00020.",
        "Touvron, H., et al. (2023). LLaMA. arXiv preprint arXiv:2302.13971.",
        "Zhu, D., et al. (2023). MiniGPT-4. arXiv preprint arXiv:2304.10592.",
        "Dai, W., et al. (2023). InstructBLIP. arXiv preprint arXiv:2305.06500."
    ]
    for r in refs:
        doc.add_paragraph(r, style='List Bullet')

    doc.save("assignment-4.docx")
    print("Assignment 4: Expanded PDF generation triggered.")

if __name__ == "__main__":
    main()
